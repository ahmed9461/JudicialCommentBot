"""Arabic RTL DOCX renderer for the final academic commentary."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .models import CommentaryDraft


class DocxRenderer:
    def __init__(self, *, font_name: str = "Arial") -> None:
        self.font_name = font_name

    def render(
        self,
        draft: CommentaryDraft,
        *,
        subject_name: str,
        output_path: Path,
        subject_slug: str | None = None,
        case_number: str | None = None,
        court_name: str | None = None,
        judgment_year: str | None = None,
        decision_number: str | None = None,
        decision_date: str | None = None,
        appeal_court_name: str | None = None,
        source_name: str | None = None,
        source_url: str | None = None,
    ) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.right_margin = Cm(2.0)
        section.left_margin = Cm(2.0)

        normal = doc.styles["Normal"]
        normal.font.name = self.font_name
        normal.font.size = Pt(14)
        r_pr = normal._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), self.font_name)
        r_fonts.set(qn("w:cs"), self.font_name)

        title = doc.add_paragraph()
        self._set_rtl(title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)
        run = title.add_run(draft.title.strip())
        run.bold = True
        run.font.name = self.font_name
        run.font.size = Pt(17)

        course = doc.add_paragraph()
        self._set_rtl(course)
        course.alignment = WD_ALIGN_PARAGRAPH.CENTER
        course.paragraph_format.space_after = Pt(6)
        course_run = course.add_run(f"المقرر: {subject_name}")
        course_run.font.name = self.font_name
        course_run.font.size = Pt(13)

        metadata_rows = [
            ("رقم القضية", case_number),
            ("المحكمة", court_name),
            ("سنة القضية", judgment_year),
            ("رقم قرار الاستئناف", decision_number),
            ("تاريخ القرار", decision_date),
            ("محكمة الاستئناف", appeal_court_name),
        ]
        metadata_rows = [(label, value) for label, value in metadata_rows if value]
        if metadata_rows:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = True
            for label, value in metadata_rows:
                cells = table.add_row().cells
                cells[0].text = str(value)
                cells[1].text = label
                for cell in cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        self._set_rtl(paragraph)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        paragraph.paragraph_format.space_after = Pt(0)
                        for item in paragraph.runs:
                            item.font.name = self.font_name
                            item.font.size = Pt(11.5)
                cells[1].paragraphs[0].runs[0].bold = True
            after_table = doc.add_paragraph()
            after_table.paragraph_format.space_after = Pt(1)

        headings = self._headings(subject_slug)
        bodies = (
            draft.facts_and_course_link,
            draft.legal_issue,
            draft.court_reasoning,
            draft.comment_and_opinion,
        )
        for heading, body in zip(headings, bodies, strict=True):
            self._add_heading(doc, heading)
            self._add_body(doc, body)

        references: list[str] = []
        for reference in draft.references:
            cleaned = reference.strip()
            if cleaned and cleaned not in references:
                references.append(cleaned)
        verified_source = ""
        if source_name and source_url:
            verified_source = f"{source_name}، المصدر الرسمي: {source_url}"
        elif source_url:
            verified_source = f"المصدر الرسمي: {source_url}"
        elif source_name:
            verified_source = source_name
        if verified_source and verified_source not in references:
            references.append(verified_source)

        if references:
            self._add_heading(doc, "المراجع")
            for reference in references:
                p = doc.add_paragraph()
                self._set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(reference)
                r.font.name = self.font_name
                r.font.size = Pt(11.5)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        footer._p.append(field)

        props = doc.core_properties
        props.author = ""
        props.last_modified_by = ""
        props.comments = ""
        doc.save(output_path)
        return output_path

    @staticmethod
    def _headings(subject_slug: str | None) -> tuple[str, str, str, str]:
        if subject_slug == "usul_al_fiqh":
            return (
                "أولاً: تلخيص وقائع القضية وربطها بالمفاهيم الأصولية",
                "ثانياً: تحديد المسألة القانونية الأصولية في الحكم وصلتها بالمقرر",
                "ثالثاً: تحليل تسبيب المحكمة في ضوء المبادئ الأصولية",
                "رابعاً: التعليق على الحكم وتبرير الرأي",
            )
        if subject_slug in {"administrative_law", "administrative_judiciary", "administrative_contracts"}:
            return (
                "أولاً: تلخيص وقائع القضية وربطها بالمفاهيم القانونية الإدارية",
                "ثانياً: تحديد المسألة القانونية وصلتها بالمقرر",
                "ثالثاً: تحليل تسبيب المحكمة في ضوء المبادئ والنصوص الإدارية",
                "رابعاً: التعليق على الحكم وتبرير الرأي",
            )
        return (
            "أولاً: تلخيص وقائع القضية وربطها بالمقرر",
            "ثانياً: تحديد المسألة القانونية وصلتها بالمقرر",
            "ثالثاً: تحليل تسبيب المحكمة",
            "رابعاً: التعليق على الحكم وتبرير الرأي",
        )

    def _add_heading(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        self._set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.bold = True
        r.font.name = self.font_name
        r.font.size = Pt(15)

    def _add_body(self, doc: Document, text: str) -> None:
        chunks = [chunk.strip() for chunk in re.split(r"\n+", text) if chunk.strip()]
        for chunk in chunks or [text.strip()]:
            p = doc.add_paragraph()
            self._set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(chunk)
            r.font.name = self.font_name
            r.font.size = Pt(14)

    @staticmethod
    def _set_rtl(paragraph) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = p_pr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)
        bidi.set(qn("w:val"), "1")
