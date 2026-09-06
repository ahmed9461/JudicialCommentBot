"""Final content gate before any Word file is sent."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from app.core.constants import FORBIDDEN_OUTPUT_MARKERS
from app.pdf.headers import labeled_case_numbers, normalize_case_number

from .models import CommentaryDraft


class CommentaryValidationError(ValueError):
    pass


_ARABIC_DIGITS = str.maketrans("٠١٢٣٤٥٦٧٨٩۰۱۲۳۴۵۶۷۸۹", "01234567890123456789")
_ARTICLE_NUMBER = re.compile(
    r"(?:المادة|المواد)\s*(?:رقم\s*)?[\(\)\[\]{}:/\-، ]{0,8}([0-9٠-٩۰-۹]{1,4})",
    re.I,
)


def _validate_text(text: str) -> None:
    lowered = text.casefold()
    for marker in FORBIDDEN_OUTPUT_MARKERS:
        if marker.casefold() in lowered:
            raise CommentaryValidationError(f"Forbidden output marker: {marker}")
    if "#" in text or "```" in text or "**" in text:
        raise CommentaryValidationError("Markdown formatting is not allowed")
    if re.search(r"(^|\s)\*(?=\S)", text):
        raise CommentaryValidationError("Markdown bullet/emphasis marker is not allowed")
    if re.search(r"\bAI\b", text, flags=re.IGNORECASE):
        raise CommentaryValidationError("AI reference is not allowed")


def _article_numbers(text: str) -> set[str]:
    return {
        match.group(1).translate(_ARABIC_DIGITS).lstrip("0") or "0"
        for match in _ARTICLE_NUMBER.finditer(text)
    }


def validate_commentary(
    draft: CommentaryDraft,
    *,
    judgment_text: str | None = None,
    expected_case_number: str | None = None,
) -> None:
    sections = (
        draft.title,
        draft.facts_and_course_link,
        draft.legal_issue,
        draft.court_reasoning,
        draft.comment_and_opinion,
        *draft.references,
    )
    for value in sections[:5]:
        if not value.strip():
            raise CommentaryValidationError("Required commentary section is empty")
    joined = "\n".join(sections)
    _validate_text(joined)

    if expected_case_number:
        expected = normalize_case_number(expected_case_number)
        foreign = [
            value
            for value in labeled_case_numbers(joined)
            if normalize_case_number(value) != expected
        ]
        if foreign:
            raise CommentaryValidationError(
                "Commentary mentions a different labelled case number: " + ", ".join(foreign)
            )

    if judgment_text is not None:
        claimed = _article_numbers(joined)
        supported = _article_numbers(judgment_text)
        unsupported = sorted(claimed - supported)
        if unsupported:
            raise CommentaryValidationError(
                "Commentary cites article number(s) not found in the verified judgment: "
                + ", ".join(unsupported)
            )


def validate_docx_file(
    path: Path,
    *,
    expected_metadata: dict[str, str | None] | None = None,
) -> None:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    if not text.strip():
        raise CommentaryValidationError("Generated DOCX contains no text")
    _validate_text(text)

    if expected_metadata:
        normalized_doc = _normalize(text)
        missing: list[str] = []
        for label, value in expected_metadata.items():
            if not value:
                continue
            if _normalize(str(value)) not in normalized_doc:
                missing.append(label)
        if missing:
            raise CommentaryValidationError(
                "Generated DOCX is missing verified judgment metadata: " + ", ".join(missing)
            )


def _normalize(value: str) -> str:
    text = value.translate(_ARABIC_DIGITS).casefold()
    text = text.replace("ـ", "")
    for source, target in (("أ", "ا"), ("إ", "ا"), ("آ", "ا"), ("ى", "ي"), ("ؤ", "و"), ("ئ", "ي")):
        text = text.replace(source, target)
    return re.sub(r"[^0-9a-z\u0600-\u06ff]+", "", text)
