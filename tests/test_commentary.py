from pathlib import Path

import pytest
from docx import Document

from app.commentary import CommentaryDraft, CommentaryValidationError, DocxRenderer, validate_commentary, validate_docx_file


def _draft(**changes) -> CommentaryDraft:
    data = {
        "title": "التعسف في استعمال الحق والتعويض عن أتعاب التقاضي",
        "facts_and_course_link": "تتلخص الوقائع في نزاع قضائي ترتب عليه طلب تعويض عن مصروفات التقاضي، وترتبط القضية بنظرية الحق وحدود استعماله المشروع في المقرر. " * 2,
        "legal_issue": "تتمثل المسألة القانونية في تحديد متى يكون اللجوء إلى القضاء استعمالاً مشروعاً للحق ومتى يتجاوز إلى التعسف الذي يرتب المسؤولية.",
        "court_reasoning": "بحثت المحكمة سبب المطالبة ومدى ثبوت التجاوز في استعمال حق التقاضي، وربطت النتيجة بالوقائع المثبتة في ملف الدعوى دون افتراض وقائع خارج الحكم. " * 2,
        "comment_and_opinion": "أؤيد النتيجة متى كان ملف الحكم لا يثبت كيدية الخصومة، لأن مجرد خسارة الدعوى السابقة لا يكفي وحده لاعتبار استعمال حق التقاضي غير مشروع. " * 2,
        "references": ["الحكم القضائي محل التعليق"],
    }
    data.update(changes)
    return CommentaryDraft(**data)


def test_validator_rejects_markdown_and_ai_mentions() -> None:
    with pytest.raises(CommentaryValidationError):
        validate_commentary(_draft(comment_and_opinion="## رأي" + " نص قانوني" * 40))
    with pytest.raises(CommentaryValidationError):
        validate_commentary(_draft(comment_and_opinion="تم إنشاء النص بواسطة AI " + "تحليل" * 40))


def test_validator_rejects_article_number_not_in_verified_judgment() -> None:
    draft = _draft(court_reasoning="استندت المحكمة إلى المادة 78 من النظام. " * 12)
    with pytest.raises(CommentaryValidationError, match="78"):
        validate_commentary(
            draft,
            judgment_text="ثبت في الحكم أن السند مستوف للشروط وفق المادة ) 87 ( من النظام.",
        )


def test_docx_renderer_is_rtl_and_contains_verified_metadata(tmp_path: Path) -> None:
    path = tmp_path / "comment.docx"
    DocxRenderer().render(
        _draft(),
        subject_name="قانون التأمين",
        subject_slug="insurance_law",
        output_path=path,
        case_number="34181612",
        court_name="المحكمة العامة بمحافظة جدة",
        judgment_year="1434",
        decision_number="35262974",
        decision_date="1435/06/02",
        appeal_court_name="محكمة الاستئناف بمنطقة مكة المكرمة",
        source_name="وزارة العدل",
        source_url="https://www.moj.gov.sa/example.pdf",
    )
    validate_docx_file(
        path,
        expected_metadata={
            "رقم القضية": "34181612",
            "المحكمة": "المحكمة العامة بمحافظة جدة",
            "سنة القضية": "1434",
            "رقم قرار الاستئناف": "35262974",
            "تاريخ القرار": "1435/06/02",
        },
    )
    doc = Document(path)
    assert "أولاً: تلخيص وقائع القضية وربطها بالمقرر" in [p.text for p in doc.paragraphs]
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "34181612" in table_text
    assert "35262974" in table_text
    assert doc.core_properties.author == ""
